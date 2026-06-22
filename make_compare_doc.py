"""生成《Lever-LM 原版对比与改进说明》Word 文档。

汇总与原版 ForJadeForest/Lever-LM 的对比问答，分三部分：
  一、我与原版代码的差别（数据使用 / anchor / 评价指标 / beam 等）
  二、多模态任务补充（少的两个任务 + 数据下载链接 + 后续代码点）
  三、Top-K 召回（原版 vs 我的 + 优化方向）
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── 页面边距 ──────────────────────────────────────
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)

# ── 中文字体（避免方框/乱码）────────────────────
def _set_cjk(style_name, ascii_font="Times New Roman", cjk="宋体", size=None):
    style = doc.styles[style_name]
    style.font.name = ascii_font
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), cjk)
    if size:
        style.font.size = Pt(size)

for s in ["Normal", "Heading 1", "Heading 2", "Heading 3", "Title"]:
    try:
        _set_cjk(s)
    except Exception:
        pass

# ── 样式辅助 ──────────────────────────────────────
def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        rpr = run._element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
        if color:
            run.font.color.rgb = RGBColor(*color)
    return p

def para(text, bold=False, size=11, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    return p

def bullet(text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    return p

def table(rows, header=None):
    cols = len(rows[0])
    t = doc.add_table(rows=len(rows) + (1 if header else 0), cols=cols)
    t.style = "Table Grid"
    if header:
        hrow = t.rows[0]
        for i, h in enumerate(header):
            hrow.cells[i].text = h
            for run in hrow.cells[i].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")
    for ri, row in enumerate(rows):
        r = t.rows[ri + (1 if header else 0)]
        for ci, val in enumerate(row):
            r.cells[ci].text = val
            for run in r.cells[ci].paragraphs[0].runs:
                run.font.size = Pt(10)
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    doc.add_paragraph()
    return t

# ════════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Lever-LM 原版对比与改进说明")
run.bold = True
run.font.size = Pt(22)
run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "黑体")

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("我的实现 vs 原版 ForJadeForest/Lever-LM（NeurIPS 2024）")
r2.font.size = Pt(13)
r2._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")

doc.add_paragraph()
para("本文档汇总三部分：(一) 代码与数据使用差异；(二) 多模态任务补充方案与数据下载；"
     "(三) Top-K 召回的原理、差异与优化。", size=10)
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 一、我与原版代码的差别
# ════════════════════════════════════════════════════════
heading("一、我与原版代码的差别", 1)

para("1.1 总览对比表", bold=True, size=13)
table([
    ["选择器模型", "GPT2 / LSTM 自回归语言模型（词表=候选池大小+3）", "PointerSelector（Cross-Attention + Pointer 打分）"],
    ["选择机制", "像写句子一样逐个生成 ICD 索引 token（自回归）", "对检索到的 K 个候选打分，贪心逐个选"],
    ["候选空间", "整个训练集是\"词表\"（闭集，池一变要重训）", "每条 query 动态检索 top-K（与池解耦，可换池）"],
    ["内容编码模型", "冻结 CLIP-ViT-L/14（图、文分别编码）", "Qwen2-VL-2B 末层隐状态 mean-pool（1536 维，图文融合）"],
    ["打分/Reward", "InfoScore = logP(y|c,x) − logP(y|x)（信息增益）", "绝对 logP(y*|S,x)（未减 baseline）"],
    ["顺序建模", "自回归，能区分 [3,9] 与 [9,3]、建模 ICD 互补", "各步独立打分，[3,9] 与 [9,3] 得分相同（学不到顺序）"],
    ["训练 Loss", "阈值过滤后所有 beam 等权 next-token CE 模仿", "RCE 奖励加权 CE，softmax(reward/τ)，τ 可学"],
    ["LVLM Backbone", "OpenFlamingo / IDEFICS / Qwen1.5", "Qwen2-VL-2B-Instruct"],
    ["任务数", "3 个（Caption / VQA / SST-2）", "1 个（VQA）"],
    ["召回采样器", "random / text-sim / img-sim / mix 四种", "仅 embedding cosine top-K 一种"],
    ["工程框架", "Hydra 配置 + 多卡 spawn", "argparse + 单卡"],
], header=["维度", "原版 Lever-LM", "我的实现"])

para("1.2 数据怎么用：什么时候用全量 44 万、什么时候用 5000", bold=True, size=13)
para("这是最容易混淆、也最该写清楚的一点。原版数据在不同环节用的是不同规模：")
para("原版数据流：", bold=True)
code_block(
    "全训练集 443,757 条（VQAv2）\n"
    "      │\n"
    "      ├─① 随机抽 5000 个 anchor（要为之生成数据的 query，sample_num=5000）\n"
    "      │\n"
    "      └─② 检索库 = 全部 443,757 条（candidate pool 是全集）\n"
    "            │\n"
    "            └─③ 为每个 anchor 从全集召回 64 个候选（cand_num=64）\n"
    "                  │\n"
    "                  └─④ 在 64 个候选上 beam search（beam_size=5, shot=2）\n"
    "                        → 输出 5 条 ICD 序列 + InfoScore"
)
para("我的数据流：", bold=True)
code_block(
    "VQAv2 训练集前 5000 条（顺序取，非随机）\n"
    "      │\n"
    "      ├─① 这 5000 条同时当 anchor（要生成数据的 query）\n"
    "      │\n"
    "      └─② 检索库 = 同一批 5000 条（anchor == 候选池）\n"
    "            │\n"
    "            └─③ 为每条 query 召回 16 个候选（K=16，排除自己）\n"
    "                  │\n"
    "                  └─④ beam search（beam_size=3, shot=2）\n"
    "                        → 输出 3 条 ICD 序列 + 绝对 logP"
)
table([
    ["要生成数据的 query 数（anchor）", "5,000（从 44 万随机抽）", "5,000（前 5000 条，非随机）"],
    ["检索库 / 候选池规模", "443,757（全集）", "5,000（前 5000）"],
    ["每条 query 召回数 K", "64", "16"],
    ["beam_size", "5", "3"],
    ["shot_num（每序列 ICD 数）", "2", "2"],
    ["anchor 与池的关系", "anchor ⊂ 全集，池≫anchor", "anchor == 池（同一批，互为示例）"],
], header=["环节", "原版用量", "我的用量"])
para("一句话：原版\"全量 44 万只作检索库（被搜的库）\"，\"5000 只是被生成数据的 anchor\"；"
     "我的实现里 5000 既是 anchor 又是检索库，没有用到全量。", bold=True)

para("1.3 anchor set（锚点集）", bold=True, size=13)
para("含义：候选池有几十万条，不可能对每条都跑昂贵的 beam search，所以先随机抽一批 query 当\"锚点\"，"
     "只为这批锚点生成 ICD 序列训练数据。原版 base_sampler.sample_anchor_set 即随机抽 anchor_sample_num 个下标并缓存。")
para("数值：anchor_sample_num = sample_num = 5000（脚本 generate_data.sh 覆盖；yaml 默认 100）。")
para("我的实现没有 anchor 概念，直接用前 5000 条全部生成。")

para("1.4 ICD 序列从哪里、怎么生成", bold=True, size=13)
para("ICD 序列不是凭空生成，而是\"从召回的候选集里、用 LVLM 当裁判、beam search 选出并排序\"：")
code_block(
    "for 每个 anchor query x:\n"
    "    candidate_set = sampler 从检索库召回的候选（原版64 / 我的16）\n"
    "    beam search：\n"
    "        shot1: 每个候选 c 拼成 [c, x] 喂 LVLM 打分 → 留 top-beam\n"
    "        shot2: 每条 beam 再配候选打分 → 留 top-beam\n"
    "    输出若干条 (ICD序列, 分数) → 作为选择器的训练监督数据"
)

para("1.5 评价指标", bold=True, size=13)
table([
    ["VQA", "VQA-accuracy（答案归一化后 exact-match）", "VQA-accuracy（同口径，归一化 exact-match）"],
    ["Caption", "CIDEr（与多条参考描述的 n-gram 重合）", "（暂无此任务）"],
    ["SST-2", "分类准确率", "（暂无此任务）"],
    ["报告主指标", "Avg:1~8 shot 准确率", "Avg:1~8 = 65.8%（500 验证样本）"],
], header=["任务", "原版指标", "我的指标"])

para("1.6 训练 Loss 与 Reward 的两个关键差异（重点改进方向）", bold=True, size=13)
bullet("Reward：原版用信息增益 logP(y|c,x)−logP(y|x)，减去了 baseline，剔除了\"答案本身好不好答\""
       "的影响，更聚焦 ICD 的边际贡献；我的是绝对 logP，受答案长度/词频影响大，也更易出现 −∞。")
bullet("Loss：原版多 beam 等权 CE 模仿；我的是 RCE 奖励加权（更先进），但因 PointerSelector 打分"
       "对顺序不敏感（[3,9] 与 [9,3] 得分相同），RCE 学不到顺序信息 → 这是当前实现最实质的能力缺失。")

para("1.7 补充：你可能没想到、但值得记录的差异", bold=True, size=13, indent=0)
bullet("候选池规模差 ≈88 倍（44.4 万 vs 5000）：这是召回质量天花板，极可能是 44.5% 样本 reward=−∞、"
       "以及准确率提升有限的主因。排查问题时应先怀疑\"池子太小 + anchor 非随机\"，再怀疑模型/Loss。")
bullet("anchor==池 的\"自指\"：query 与候选高度同源、分布一致，多样性受限，易过拟合该小段分布。")
bullet("\"前 5000 条\"非随机带来的采样偏差：开头可能集中某些图/题型（如大量 yes/no），任务多样性不足。")
bullet("规模仅论文约 0.5%：val/loss 卡在随机下界 log(16)×2≈5.55 附近，更可能是数据规模所致，而非方法无效。")
bullet("检索引擎：原版用 FAISS（IndexFlatIP，归一化后内积=余弦）；我的用 torch.topk（规模小够用）。")
bullet("特殊 token 与词表：原版 GPT2 词表= 池大小+3（BOS/EOS/QUERY），换池子要重训；我的 PointerSelector 与池解耦，换池/换任务不改模型——这是我方架构的优势。")

# ════════════════════════════════════════════════════════
# 二、多模态任务补充
# ════════════════════════════════════════════════════════
heading("二、多模态任务补充", 1)

para("现状：我只实现了 VQA 一个任务，原版有 3 个，我缺 2 个：Image Captioning 与 SST-2。"
     "补齐后可与原版做公平对比，并证明方法的通用性（覆盖\"开放生成 / 短答 / 封闭分类\"三种范式）。")

para("2.1 原版三个任务介绍（举例）", bold=True, size=13)
para("① Image Captioning（图像描述，COCO）—— 缺", bold=True)
para("输入只有一张图、无问题；输出一句开放式描述；评测用 CIDEr。")
code_block(
    "ICD_1: [猫趴沙发图] → \"a cat lying on a sofa\"\n"
    "ICD_2: [自行车图]   → \"a bicycle parked on the street\"\n"
    "Query: [披萨图]     → 生成 \"a pizza on a plate\""
)
para("② VQA（视觉问答，VQAv2）—— 已有", bold=True)
code_block("Query: [蓝天图] \"What is in the sky?\" → 生成 \"clouds\"")
para("③ SST-2（情感分类，纯文本）—— 缺", bold=True)
para("输入一句话、无图；输出 2 分类标签（positive/negative）；评测用分类准确率。"
     "Qwen2-VL 也能处理纯文本，prompt 里不放图即可。")
code_block(
    "ICD_1: \"the movie was boring and too long\" → negative\n"
    "ICD_2: \"a delightful and touching story\"   → positive\n"
    "Query: \"the acting saved an otherwise dull plot\" → 生成 positive"
)

para("2.2 需要下载哪些数据（含链接）", bold=True, size=13)
para("① COCO Caption（用于 Image Captioning，对应原版 coco2017，train≈118,287 张）", bold=True)
bullet("训练图片：http://images.cocodataset.org/zips/train2017.zip")
bullet("验证图片：http://images.cocodataset.org/zips/val2017.zip")
bullet("标注（含 caption）：http://images.cocodataset.org/annotations/annotations_trainval2017.zip")
bullet("HuggingFace 镜像（可选，省去手动下载）：https://huggingface.co/datasets/HuggingFaceM4/COCO")
bullet("Karpathy split（论文常用划分，可选）：https://huggingface.co/datasets/yerevann/coco-karpathy")

para("② SST-2（用于情感分类，纯文本，train≈67k）", bold=True)
bullet("HuggingFace 官方：https://huggingface.co/datasets/stanfordnlp/sst2")
bullet("GLUE 版本（含标准 dev 集）：https://huggingface.co/datasets/nyu-mll/glue（config 选 sst2）")
bullet("原始下载：https://dl.fbaipublicfiles.com/glue/data/SST-2.zip")

para("③ VQAv2（已具备，列出备查）", bold=True)
bullet("HuggingFace：https://huggingface.co/datasets/lmms-lab/VQAv2（你已用本地 parquet 快照）")

para("2.3 后续写代码要改的三处（数据下载后再做，先了解）", bold=True, size=13)
bullet("数据加载：为 COCO、SST-2 各写 loader，统一成 {id, image(可空), text/question, answer} 格式。")
bullet("Reward 打分：Caption 改成\"先生成再算 CIDEr\"；SST-2 改成\"比较各标签 log-prob 取 argmax\"；"
       "VQA 维持现状。借此机会顺带把\"增益式打分 + 修 −∞\"一起做掉。")
bullet("评测指标：Caption→CIDEr、VQA→acc、SST-2→分类准确率；把现在硬编码的 VQA exact-match 抽成 per-task 函数。")
bullet("PointerSelector 本体几乎不动（只吃 embedding），换任务不改模型——这是相对原版的优势。")

# ════════════════════════════════════════════════════════
# 三、Top-K 召回
# ════════════════════════════════════════════════════════
heading("三、Top-K 召回", 1)

para("3.1 它是什么", bold=True, size=13)
para("召回 = 第一阶段粗筛。池子有几千~几十万条，逐条用 LVLM 打分太贵，所以先用一个便宜的相似度"
     "把池子缩到 K 条，后续 beam search / PointerSelector 只在这 K 条上工作。它决定了性能天花板——"
     "好 ICD 没进 top-K，后面再聪明也选不到。")

para("3.2 原版怎么做（模型 + 方法）", bold=True, size=13)
bullet("编码模型：冻结 CLIP-ViT-L/14，图、文分别编码并 L2 归一化。")
bullet("检索引擎：FAISS IndexFlatIP（归一化向量内积 = 余弦相似度），取 top_k=K+1 再去掉自己。")
bullet("检索库：整个训练集（VQA 全集 443,757；Caption 全集 118,287）。")
bullet("四种采样器：random / text-sim（按 question 等文本字段）/ img-sim（按图）/ mix（按比例混合）。")
bullet("mix 比例（原版主推）：RandSampler 0.5 + TextSimSampler 0.25 + ImgSimSampler 0.25，K=64。")
bullet("结果会缓存（feature_cache），不必每次重算。")

para("3.3 我怎么做", bold=True, size=13)
bullet("编码模型：Qwen2-VL-2B 本身，把 (图+问题) 一起喂入取末层隐状态 mean-pool 成 1536 维（图文融合在一个向量里）。")
bullet("检索：query 向量与池向量 L2 归一化后算余弦，torch.topk 取 top-K，排除自己。")
bullet("检索库：候选池 5000 条；K=16；只有这一种 cosine 策略，无 mix、无多样性控制。")

para("3.4 原版 vs 我的（对比）", bold=True, size=13)
table([
    ["编码模型", "CLIP-ViT-L/14（图/文分离）", "Qwen2-VL-2B（图文融合一个向量）"],
    ["检索库规模", "443,757（全集）", "5,000"],
    ["召回数 K", "64", "16"],
    ["召回策略", "random/text/img/mix 四选一+混合", "仅 cosine top-K"],
    ["能否按问题单独召回", "能（text-sim 按 question）", "不能（图文已揉成一个向量，图易主导）"],
    ["检索引擎", "FAISS", "torch.topk"],
], header=["维度", "原版", "我的"])

para("3.5 优化方向（按性价比排序）", bold=True, size=13)
bullet("A. 先量化召回质量（强烈建议第一步，零成本）：统计 beam 选中的好 ICD 落在 top-K 的比例"
       "（Recall@K）。若很低，说明瓶颈在召回不在选择器，可直接写进报告。")
bullet("B. 引入混合召回（对标原版 mix）：图相似 + 问题文本相似 + 少量随机，提升多样性与覆盖；"
       "VQA 尤其要加\"按 question 文本召回\"，避免图主导、问题信号被淹没。")
bullet("C. 召回向量与选择器向量解耦：召回用检索更强的向量（如 CLIP/对比学习向量），选择器仍用 Qwen 向量。")
bullet("D. 增大 K（16→32/64）：抬高天花板、贴近论文设置；代价是数据生成打分次数上升（算力权衡）。")
bullet("E. 多样性感知召回（MMR）：相似度里减去与已召回候选的冗余度，避免 K 个里一半是近重复。")
bullet("F. 学习式检索（偏研究）：用 beam 数据反向训一个双塔 retriever + 难负样本，让召回本身可学。")
bullet("补充 G. 把检索从 torch.topk 换成 FAISS：当池子扩到几十万时，torch.topk 会变慢，FAISS 必要。")
bullet("补充 H. query 端与候选端可用不同字段/模态构造（如 VQA 用问题文本召回、Caption 用图召回），按任务配置化。")
para("建议优先级：先 A（量化）→ B（混合召回，尤其问题文本召回）→ D（适度加大 K）；C/E/F/G 视报告深度再加。",
     bold=True)

# ════════════════════════════════════════════════════════
# 结尾
# ════════════════════════════════════════════════════════
doc.add_paragraph()
para("附：本文档基于对原版仓库 ForJadeForest/Lever-LM 源码与我的实现的逐文件对比整理。"
     "思考题（embedding 表示 / VAE / 128×768→1×768 等）待后续补充。", size=9)

import os
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Lever-LM_原版对比与改进说明.docx")
doc.save(out)
print("Saved:", out)
