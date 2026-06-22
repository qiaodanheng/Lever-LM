"""记录 Base 模型半套实验：推理用 Base，Step1/Selector 仍来自 Instruct 训练数据。"""
from datetime import date
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
import json

TODAY = date.today().strftime("%Y-%m-%d")


def cjk(run, name="宋体"):
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)


def h(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        cjk(r, "黑体")


def p(doc, text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    cjk(run)


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.size = Pt(11)
    cjk(run)


def table(doc, header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Table Grid"
    for i, x in enumerate(header):
        t.rows[0].cells[i].text = x
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
            cjk(r, "黑体")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for r in t.rows[ri + 1].cells[ci].paragraphs[0].runs:
                cjk(r)


def acc(path):
    return json.load(open(path))["accuracy"] * 100


def sweep_avg(path):
    return json.load(open(path))["avg_1_8"] * 100


doc = Document()
sec = doc.sections[0]
sec.top_margin = sec.bottom_margin = Cm(2.5)
sec.left_margin = Cm(3.0)
sec.right_margin = Cm(2.5)

h(doc, f"Lever-LM Base 半套实验记录（{TODAY}）", 0)
p(doc, "实验类型：Qwen2-VL-2B（Base）推理 + Instruct 训练的 Step1/Selector")
p(doc, f"日期：{TODAY}  ·  目录：/home/jiyi/lizhiheng/Lever-LM")

h(doc, "一、实验设定（半套 Base 的含义）", 1)
p(doc,
  "本实验用于初步对比 Base 与 Instruct 在 VQA ICL 上的差异，"
  "并非从 Step1 重做的「干净 Base 全流程」。"
  "仅将下列环节换成 Base 模型，其余仍沿用 Instruct 产物：")

table(doc, ["环节", "Base 半套实验", "仍用 Instruct 的部分"], [
    ["Step0 池 embed", "vqav2_pool_v2_base / vqav2_pool_full_base", "—"],
    ["Step1 Beam 数据", "—", "vqav2_train_beams_v2.json"],
    ["Step2 Selector", "—", "checkpoints/lever_lm_v2_poolneg/epoch003…ckpt"],
    ["Step3 query embed", "Base", "—"],
    ["Step3 检索池", "Base 池向量", "—"],
    ["Step3 Lever 选 ICD", "—", "Instruct 上训练的 Selector"],
    ["Step3 生成答案", "Base", "—"],
])

h(doc, "二、动机说明", 1)
bullet(doc, "怀疑 Instruct 在预训练/对齐阶段见过 VQA 类数据，导致 0-shot 虚高、ICL 反而低于 0-shot。")
bullet(doc, "换 Base 模型做对照：若 0-shot 明显下降，说明 Instruct 在 VQA 上确有额外优势（可能来自对齐或数据重叠，本实验不能单独证明「泄露」）。")
bullet(doc, "本记录归档半套结果；完整 Base 全流程（Step1→训练→ICL）见 scripts/run_base_full_pipeline.sh。")

h(doc, "三、配置", 1)
table(doc, ["参数", "值"], [
    ["大模型（推理）", "/home/jiyi/.cache/modelscope/qwen/Qwen2-VL-2B"],
    ["Selector ckpt", "checkpoints/lever_lm_v2_poolneg/epoch003-valloss7.0768.ckpt"],
    ["20k 池", "data/vqav2_pool_v2_base.jsonl + .pt"],
    ["443k 池", "data/vqav2_pool_full_base.jsonl + .pt"],
    ["K", "32"],
    ["shot", "2（sweep 1~8）"],
    ["评测", "VQAv2 validation 前 5000，eval_protocol=unified"],
])

h(doc, "四、实验结果（Base 半套）", 1)

h(doc, "4.1 20k 池，val 5000，shot=2", 2)
table(doc, ["方法", "Acc"], [
    ["Lever", f"{acc('results/icl_base_20k_lever_unified_n5000.json'):.2f}%"],
    ["Random", f"{acc('results/icl_base_20k_random_unified_n5000.json'):.2f}%"],
    ["0-shot", f"{acc('results/icl_base_20k_zeroshot_unified_n5000.json'):.2f}%"],
    ["Lever−Random", f"{acc('results/icl_base_20k_lever_unified_n5000.json')-acc('results/icl_base_20k_random_unified_n5000.json'):+.2f}%"],
])

h(doc, "4.2 20k 池，shot sweep Avg:1~8", 2)
table(doc, ["方法", "Avg:1~8"], [
    ["Lever", f"{sweep_avg('results/icl_base_20k_lever_unified_n5000.sweep.sweep_summary.json'):.2f}%"],
    ["Random", f"{sweep_avg('results/icl_base_20k_random_unified_n5000.sweep.sweep_summary.json'):.2f}%"],
    ["0-shot", f"{sweep_avg('results/icl_base_20k_zeroshot_unified_n5000.sweep.sweep_summary.json'):.2f}%"],
])

h(doc, "4.3 443k 池，val 5000，shot=2，K=32", 2)
table(doc, ["方法", "Acc"], [
    ["Lever", f"{acc('results/icl_base_fullpool_lever_unified_n5000_k32.json'):.2f}%"],
    ["Random", f"{acc('results/icl_base_fullpool_random_unified_n5000_k32.json'):.2f}%"],
    ["0-shot", f"{acc('results/icl_base_fullpool_zeroshot_unified_n5000_k32.json'):.2f}%"],
    ["Lever−Random", f"{acc('results/icl_base_fullpool_lever_unified_n5000_k32.json')-acc('results/icl_base_fullpool_random_unified_n5000_k32.json'):+.2f}%"],
])

h(doc, "五、与 Instruct 对照", 1)
table(doc, ["设置", "Instruct Lever", "Base 半套 Lever", "Instruct 0-shot", "Base 0-shot"], [
    ["20k shot=2", "63.38%", f"{acc('results/icl_base_20k_lever_unified_n5000.json'):.2f}%",
     "70.54%", f"{acc('results/icl_base_20k_zeroshot_unified_n5000.json'):.2f}%"],
    ["443k shot=2", "63.44%", f"{acc('results/icl_base_fullpool_lever_unified_n5000_k32.json'):.2f}%",
     "70.54%", f"{acc('results/icl_base_fullpool_zeroshot_unified_n5000_k32.json'):.2f}%"],
    ["20k Avg:1~8", "62.01%", f"{sweep_avg('results/icl_base_20k_lever_unified_n5000.sweep.sweep_summary.json'):.2f}%",
     "70.54%", "55.66%"],
])

h(doc, "六、结论与局限", 1)
bullet(doc, "Base 0-shot（55.66%）远低于 Instruct（70.54%），支持「Instruct 在 VQA 短答上显著更强」。")
bullet(doc, "Base 上 ICL（Lever ~53.7%）仍略高于 Random（~52.8~53.4%），但二者均低于 0-shot。")
bullet(doc, "Lever 使用 Instruct 训练的 Selector，在 Base 向量空间上存在分布错位，Lever 增益被压缩（Lever−Random 仅 +0.5~1%）。")
bullet(doc, "本实验不能严谨证明 VQA 数据泄露；要验证需做 Base 全流程或 val 重叠检测。")

h(doc, "七、结果文件索引", 1)
for f in [
    "results/icl_base_20k_lever_unified_n5000.json",
    "results/icl_base_20k_random_unified_n5000.json",
    "results/icl_base_20k_zeroshot_unified_n5000.json",
    "results/icl_base_20k_*_unified_n5000.sweep*",
    "results/icl_base_fullpool_*_unified_n5000_k32.json",
    "data/vqav2_pool_v2_base.jsonl / .pt",
    "data/vqav2_pool_full_base.jsonl / .pt",
    "data/icl_base_qwen_pipeline.log",
]:
    bullet(doc, f)

out = f"/home/jiyi/lizhiheng/Lever-LM/Lever-LM_Base半套实验记录_{TODAY}.docx"
doc.save(out)
print("Saved →", out)
