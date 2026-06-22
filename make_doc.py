"""生成项目说明 Word 文档"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 页面边距 ──────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.0)
section.right_margin  = Cm(2.5)

# ── 样式辅助函数 ──────────────────────────────────
def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def para(text, bold=False, size=11, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def code_block(text):
    """灰底等宽字体代码块"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    # 设置底纹
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return p

def table_2col(rows, header=None):
    cols = 2
    t = doc.add_table(rows=len(rows) + (1 if header else 0), cols=cols)
    t.style = 'Table Grid'
    if header:
        hrow = t.rows[0]
        for i, h in enumerate(header):
            hrow.cells[i].text = h
            for run in hrow.cells[i].paragraphs[0].runs:
                run.bold = True
    for ri, row in enumerate(rows):
        r = t.rows[ri + (1 if header else 0)]
        for ci, val in enumerate(row):
            r.cells[ci].text = val
    doc.add_paragraph()

def table_3col(rows, header=None):
    t = doc.add_table(rows=len(rows) + (1 if header else 0), cols=3)
    t.style = 'Table Grid'
    if header:
        hrow = t.rows[0]
        for i, h in enumerate(header):
            hrow.cells[i].text = h
            for run in hrow.cells[i].paragraphs[0].runs:
                run.bold = True
    for ri, row in enumerate(rows):
        r = t.rows[ri + (1 if header else 0)]
        for ci, val in enumerate(row):
            r.cells[ci].text = val
    doc.add_paragraph()

def divider():
    doc.add_paragraph('─' * 60)

# ════════════════════════════════════════════════════════
# 封面
# ════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Lever-LM 项目说明文档')
run.bold = True
run.font.size = Pt(20)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run('基于 Qwen2-VL-2B 的 In-Context Learning 示例选择').font.size = Pt(13)

doc.add_paragraph()
doc.add_paragraph()

# ════════════════════════════════════════════════════════
# 1. 任务目标
# ════════════════════════════════════════════════════════
heading('一、任务目标', 1)

para('任务：VQA（Visual Question Answering，视觉问答）', bold=True)
para('给定一张图片和一个自然语言问题，让模型输出文字答案。')
code_block('输入：[图片]  +  "What color is the car?"\n输出："red"')
para('数据集：VQAv2（基于 Microsoft COCO）')
table_2col([
    ['训练集', '443,757 条问答对'],
    ['验证集', '214,354 条问答对'],
], header=['数据集拆分', '样本数量'])

# ════════════════════════════════════════════════════════
# 2. 核心思路
# ════════════════════════════════════════════════════════
heading('二、核心思路：用示例帮助大模型', 1)

para('大模型不需要 fine-tune，只需在问题前面拼几个"示例"（ICD），回答质量就会大幅提升，这种方法叫 In-Context Learning（ICL）。')

code_block(
    '示例1：[狗的图]  Q: "What animal?"  A: "dog"\n'
    '示例2：[红车图]  Q: "What color?"   A: "red"\n'
    '────────────────────────────────────────────\n'
    '测试：[蓝天图]  Q: "What is in the sky?"  A: ???'
)

para('关键问题：示例选哪些？顺序怎么排？选不好效果差很多。')
para('Lever-LM（NeurIPS 2024）的解决方案：', bold=True)
para('训练一个小模型（PointerSelector，~527K 参数），专门负责"从候选里挑最好的示例顺序"，让大模型回答得更准。', indent=1)

# ════════════════════════════════════════════════════════
# 3. 两个模型的分工
# ════════════════════════════════════════════════════════
heading('三、两个模型的分工', 1)

table_3col([
    ['Qwen2-VL-2B', '2B 参数', '理解图像，给 ICD 序列打分，生成最终答案'],
    ['PointerSelector', '527K 参数', '学习"选哪些示例、什么顺序"'],
], header=['模型', '参数量', '作用'])

table_3col([
    ['Qwen2-VL-2B', '不训练（frozen）', '数据生成打分 + 最终推理'],
    ['PointerSelector', '训练（RCE Loss）', '数据生成完成后训练，推理时使用'],
], header=['模型', '是否训练', '使用时机'])

# ════════════════════════════════════════════════════════
# 4. 整体流程
# ════════════════════════════════════════════════════════
heading('四、整体流程（三大步）', 1)

table_3col([
    ['Step 1', '数据生成', '用 Qwen2-VL 打分，beam search 找最优示例序列'],
    ['Step 2', '训练 PointerSelector', '用 Step 1 的数据训练小模型'],
    ['Step 3', '评估', '用 PointerSelector 选示例，测 VQA 准确率'],
], header=['步骤', '名称', '说明'])

table_3col([
    ['Step 1', '~7 小时', '正在运行中'],
    ['Step 2', '~5 分钟', '等 Step 1 完成后运行'],
    ['Step 3', '~1 小时', '等 Step 2 完成后运行'],
], header=['步骤', '预计耗时', '状态'])

# ════════════════════════════════════════════════════════
# 5. Step 1 详解
# ════════════════════════════════════════════════════════
heading('五、Step 1 数据生成详解', 1)

heading('5.1 运行参数', 2)
table_2col([
    ['anchor 样本数', '5000 条（从训练集随机取）'],
    ['候选池大小 (K)', '16（每个 query 检索 16 个候选）'],
    ['beam size',     '3（保留 top-3 条序列）'],
    ['shot_num',      '2（每条序列含 2 个 ICD）'],
    ['打分模型',       'Qwen2-VL-2B-Instruct（本地）'],
], header=['参数', '值'])

heading('5.2 每条 Query 的处理流程', 2)
code_block(
    '① 用 Qwen2-VL 提取 query 的 embedding（2048 维向量）\n'
    '② 余弦相似度检索候选池，取 top-16 个候选\n'
    '③ Beam Search（2步）：\n'
    '   Step 0（选第1个ICD）：\n'
    '     对 16 个候选各打一次分\n'
    '     分数 = log P(正确答案 | [候选i] + [query])\n'
    '     保留分数最高的 top-3 作为 beam\n'
    '   Step 1（选第2个ICD）：\n'
    '     对每个 beam，对剩余15个候选各打分\n'
    '     3 × 15 = 45 次打分\n'
    '     保留总分最高的 top-3 完整序列\n'
    '④ 输出：3条完整的 2-shot ICD 序列 + 各自 reward 分数'
)

heading('5.3 打分公式（论文 Eq. 3）', 2)
para('score = log P(y* | S^K, x)')
para('含义：模型在看到 ICD 序列 S^K 和 query x 后，生成正确答案 y* 的对数概率。分数越高 → 这组示例对模型越有帮助。')

heading('5.4 输出格式', 2)
code_block(
    '{\n'
    '  "query_id": "123",\n'
    '  "query_emb": [0.12, -0.34, ...],      // query 的 2048 维 embedding\n'
    '  "cand_embs": [[...], [...], ...],      // 16 个候选的 embedding\n'
    '  "beam_labels": [[2,7], [2,11], [5,7]],// 3条序列（候选局部索引）\n'
    '  "beam_rewards": [-12.3, -13.1, -14.5] // 对应 reward 分数\n'
    '}'
)

heading('5.5 总打分次数', 2)
para('每条 query：16 + 3×15 = 61 次 Qwen-VL 前向传播')
para('5000 条 query 合计：5000 × 61 = 305,000 次前向传播')
para('预计耗时：~7 小时（RTX 5090，已实现批量打分优化，1.4× 加速）')

# ════════════════════════════════════════════════════════
# 6. Step 2 详解
# ════════════════════════════════════════════════════════
heading('六、Step 2 训练 PointerSelector 详解', 1)

heading('6.1 模型结构', 2)
code_block(
    '输入：\n'
    '  query_emb : [1,  2048]   ← query 的 embedding\n'
    '  cand_embs : [16, 2048]   ← 16 个候选的 embedding\n'
    '\n'
    '内部：\n'
    '  Linear(2048 → 512)\n'
    '  TransformerEncoder（2层，8头 attention）\n'
    '    → query 和候选一起做 self-attention\n'
    '    → 学习"哪些候选组合在一起最好"\n'
    '  Pointer head：每步输出 K 个候选的得分分布\n'
    '\n'
    '输出：\n'
    '  每步对 K 个候选的评分（greedy / beam search 时逐步选）'
)

heading('6.2 损失函数：RCE（Ranked Cross-Entropy）', 2)
para('训练数据来自 Step 1：每条 query 有 3 条序列，按 reward 排序：seq1 > seq2 > seq3')
para('RCE 的直觉：')
para('• 第1步选择：告诉模型，"最好序列的第1个ICD"应该得最高分', indent=1)
para('• 第2步选择：给定第1步选择，"最好序列的第2个ICD"应该得最高分', indent=1)
para('• 排名越好的序列，对应的 loss 权重越大', indent=1)

heading('6.3 训练配置', 2)
table_2col([
    ['训练样本数',    '5000 × 3 = 15000 条'],
    ['优化器',       'AdamW，lr = 1e-4'],
    ['Epochs',       '20'],
    ['Batch size',   '128'],
    ['学习率调度',   'Cosine scheduler，5% warmup'],
    ['预计耗时',     '~5 分钟（RTX 5090）'],
], header=['配置项', '值'])

# ════════════════════════════════════════════════════════
# 7. Step 3 详解
# ════════════════════════════════════════════════════════
heading('七、Step 3 评估详解', 1)

heading('7.1 推理流程', 2)
code_block(
    '对验证集中每条 query：\n'
    '① 用 Qwen2-VL 提取 query embedding\n'
    '② 余弦相似度从训练集候选池检索 top-K\n'
    '③ 用训练好的 PointerSelector 选择最优 ICD 序列（毫秒级）\n'
    '④ 拼接 [ICD序列] + [query] 输入 Qwen2-VL\n'
    '   让 Qwen2-VL generate 出答案\n'
    '⑤ 对比生成答案和 ground truth，计算准确率'
)

heading('7.2 评估指标', 2)
para('测试 shot_num = 1, 2, 3, 4, 5, 6, 7, 8 共 8 种配置：')
table_2col([
    ['Avg:1~2', '插值能力（训练时用2-shot，测1和2的平均准确率）'],
    ['Avg:3~8', '外推能力（没训练过的shot数，测泛化性）'],
    ['Avg:1~8', '整体性能（主要对比指标）'],
], header=['指标', '含义'])

# ════════════════════════════════════════════════════════
# 8. 代码文件说明
# ════════════════════════════════════════════════════════
heading('八、代码文件说明', 1)

table_3col([
    ['qwen_vl_scorer.py',  '核心算法', '用 Qwen2-VL 给 ICD 序列打分（Eq. 3），支持批量打分加速'],
    ['generate_data.py',   '核心算法', 'Beam Search 搜索最优 ICD 序列，生成训练数据'],
    ['pointer_selector.py','核心算法', 'PointerSelector 模型定义（小 Transformer）'],
    ['lever_lm_module.py', '核心算法', 'RCE 损失函数 + PyTorch Lightning 训练模块'],
    ['dataset.py',         '工程代码', '加载 VQAv2 数据集，封装成训练用格式'],
    ['train.py',           '工程代码', '训练入口脚本'],
    ['icl_inference.py',   '工程代码', '评估脚本（测试 shot 1~8 的 VQA 准确率）'],
], header=['文件名', '类型', '作用'])

# ════════════════════════════════════════════════════════
# 9. 与计算机视觉概念的对照
# ════════════════════════════════════════════════════════
heading('九、与你熟悉的 CV 概念对照', 1)

table_2col([
    ['特征图（Feature Map）',      '视觉 Token（图像 patch 编码后的序列）'],
    ['FPN / Neck',                 'Cross-Attention Fusion（视觉与语言融合）'],
    ['训练集标注（GT box/mask）',  'Beam Search 生成的 ICD 序列 + reward（软标签）'],
    ['分类 Head 输出概率',         'log P(answer | context)（生成答案的对数概率）'],
    ['NMS 保留 top-k 框',          'Beam Search 保留 top-k 序列'],
    ['mAP / mIoU',                 'VQA Accuracy（Avg:1~8）'],
    ['预训练 ResNet/ViT',           'Qwen2-VL（预训练 VLM，不 fine-tune）'],
    ['轻量化检测头',               'PointerSelector（小 Transformer，只训练这个）'],
], header=['你熟悉的 CV 概念', '本项目对应概念'])

# ════════════════════════════════════════════════════════
# 10. 学习路线
# ════════════════════════════════════════════════════════
heading('十、学习路线建议', 1)

para('你已有深度学习基础（CNN、检测、分割），需要补充以下方向：')

heading('10.1 Transformer 架构（最重要）', 2)
para('你已接触过（DETR、SegFormer 等），需要深入理解：')
para('• Self-Attention / Cross-Attention 计算过程', indent=1)
para('• Decoder-only 架构（GPT 系列）vs Encoder-Decoder（T5）', indent=1)
para('• 推荐：Andrej Karpathy《Let\'s build GPT》视频', indent=1)

heading('10.2 大语言模型基础', 2)
para('• Token 和 Tokenization（类比像素和 patch）', indent=1)
para('• Autoregressive 生成（逐 token 输出，类比逐步 decode）', indent=1)
para('• log-prob、temperature、beam search', indent=1)
para('• 推荐：李宏毅 2023/2024 大模型课程', indent=1)

heading('10.3 视觉-语言模型（VLM）', 2)
para('• 图像如何变成 token 喂给 LLM（ViT patch → visual tokens）', indent=1)
para('  类比：你熟悉的 feature map → 变成序列送入 Transformer', indent=1)
para('• 代表模型：LLaVA、Qwen-VL、InternVL、GPT-4V', indent=1)
para('• 推荐：读 LLaVA 论文（简单易懂）', indent=1)

heading('10.4 In-Context Learning（ICL）', 2)
para('• 本项目的核心概念', indent=1)
para('• 推荐论文：GPT-3（Brown et al. 2020）→ RICES → Lever-LM', indent=1)

heading('10.5 工程工具', 2)
para('• HuggingFace transformers 库（类比 torchvision）', indent=1)
para('• HuggingFace datasets 库', indent=1)
para('• PyTorch Lightning（结构化训练循环）', indent=1)

heading('10.6 推荐学习顺序', 2)
code_block(
    'Transformer 原理  →  LLM 基础  →  HuggingFace 实践  →  VLM  →  ICL\n'
    '    (1~2周)           (1周)            (1周)           (1周)    (1周)'
)

# ════════════════════════════════════════════════════════
# 保存
# ════════════════════════════════════════════════════════
out = '/home/jiyi/lizhiheng/Lever-LM/Lever-LM项目说明.docx'
doc.save(out)
print(f'已保存：{out}')
