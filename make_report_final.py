"""生成 Lever-LM v2 最终实验总结 Word 文档（含与原版对比、问题与解决方案）"""
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

doc = Document()
section = doc.sections[0]
section.top_margin = section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.0)
section.right_margin = Cm(2.5)


def cjk(run, name="宋体"):
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), name)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        cjk(r, "黑体")
    return p


def p(text, bold=False):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    cjk(run)
    return para


def bullet(text):
    para = doc.add_paragraph(style="List Bullet")
    run = para.add_run(text)
    run.font.size = Pt(11)
    cjk(run)


def table(header, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Table Grid"
    for i, x in enumerate(header):
        t.rows[0].cells[i].text = x
        for r in t.rows[0].cells[i].paragraphs[0].runs:
            r.bold = True
            cjk(r, "黑体")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
            for r in t.rows[ri + 1].cells[ci].paragraphs[0].runs:
                cjk(r)


# ══════════════════════════════════════════════════════════════
h("Lever-LM（Qwen2-VL 版）实验总结报告", 0)
p("基于 Qwen2-VL-2B 的多模态 In-Context Learning 示例选择 · 与原版 Lever-LM 对比")
p("日期：2026-06-06 · 代码目录：/home/jiyi/lizhiheng/Lever-LM")

# ── 一、项目目标与整体流程 ──
h("一、项目目标与整体流程", 1)
p(
    "在「不微调 Qwen2-VL-2B」的前提下，训练小型 PointerSelector，"
    "为每个 query 从候选池中有序选出 in-context 示例（ICD），使大模型 VQA 等任务表现更好。"
    "核心诉求：① 解决 reward=−∞；② 解决 [3,9] 与 [9,3] 顺序不可区分；③ 对齐三任务（VQA/Caption/SST-2）。"
)

h("1.1 四步流程", 2)
table(["步骤", "内容", "本实验状态"], [
    ["Step 0 候选池", "train 中 (图,问,答) + Qwen 向量，供检索", "2 万 → 已扩至 443,757 全 train"],
    ["Step 1 数据生成", "5000 anchor → 检索 top-K → beam → InfoScore", "三任务完成，0% −∞"],
    ["Step 2 训练 Selector", "PointerSelector + RCE + 池负样本 InfoNCE", "三任务 ckpt 完成"],
    ["Step 3 下游 ICL", "检索 → 选择 → Qwen 生成 → VQA acc", "VQA val5000 完成；Caption/SST-2 未做"],
])

h("1.2 ICL 推理链路（Step 3）", 2)
bullet("检索：query 向量与池内向量算余弦相似度，取 top-K=32（菜单）")
bullet("选择：PointerSelector 从 32 个中 greedy 选 shot_num=2 个并排顺序")
bullet("推理：_build_messages 拼成多轮对话（2 个 ICD 各含图+问+assistant 答案 + 当前题），Qwen 生成")
bullet("评测：unified 协议（三方法均短答后缀 + 答案抽取 + strict 匹配）")

# ── 二、相对 v1 的改进 ──
h("二、相对 v1 的工程与创新", 1)
table(["维度", "v1（旧）", "v2（新）"], [
    ["Reward", "子串匹配，44.5% −∞", "Teacher-forcing + InfoScore，0% −∞"],
    ["顺序", "[3,9]=[9,3] 不可区分", "自回归 PointerSelector，排列对 acc 0.67~0.74"],
    ["数据", "前 5000 顺序取，anchor=池", "随机 5000 anchor + 解耦池"],
    ["任务", "仅 VQA", "VQA + Caption + SST-2"],
    ["嵌入维", "768（错误）", "1536（Qwen2-VL）"],
    ["训练", "仅 RCE", "RCE + 池负样本 InfoNCE"],
])

# ── 三、与原版 Lever-LM 论文的区别 ──
h("三、与原版 Lever-LM（NeurIPS 2024）的区别", 1)

h("3.1 架构与实验条件", 2)
table(["维度", "原版 Lever-LM", "本实验 v2"], [
    ["大模型", "OpenFlamingo / IDEFICS 9B", "Qwen2-VL-2B"],
    ["小模型", "67M Transformer/LSTM 自回归", "PointerSelector 指针网络"],
    ["特征/检索", "CLIP + 全 train FAISS/相似度", "Qwen hidden 1536 维 + top-K"],
    ["候选池", "全 train ~44 万", "先 2 万，后 443,757 全量"],
    ["Step1 anchor", "随机 5000", "随机 5000（已对齐）"],
    ["Step1 子支持", "随机 64 / 相似检索", "top-K=32，beam=3"],
    ["Reward", "InfoScore", "InfoScore（已对齐）"],
    ["ICL 形式", "9B 原生交错 ICL", "Qwen 多图 chat（每 ICD 独立图）"],
])

h("3.2 评测协议：论文无 0-shot baseline", 2)
p(
    "重要说明：原版 Lever-LM 论文 Table 1 仅报告 shot=1~8 的 ICL 方法对比"
    "（Lever-LM vs RS 随机例题 vs SIIR/SITR/STTR 相似度检索等），"
    "主表未设置 0-shot（不加任何 ICD）baseline。"
)
p(
    "合理推测：在 9B LVLM 上 0-shot 基线相对较低，ICL 有提升空间；"
    "论文聚焦「在同样使用 ICL 的前提下，谁选例题更好」。"
    "本实验额外评测 0-shot 是为诊断；"
    "发现 Qwen2-VL-2B 0-shot 已达 70.5%，高于 ICL（63.4%），"
    "这与论文实验条件不同，不宜直接对比绝对分数。"
)

h("3.3 原版论文 ICL 效果（供参照，非直接可比）", 2)
table(["任务 / 模型", "Lever-LM Avg:1~8", "RS（随机 ICD）", "Lever−RS"], [
    ["VQA · OpenFlamingo-9B", "51.31%", "47.94%", "+3.37%"],
    ["VQA · IDEFICS-9B", "54.44%", "53.54%", "+0.90%"],
    ["Caption CIDEr · OpenFlamingo", "92.45", "88.48", "+3.97"],
])
p("论文结论：在「均使用 ICL」的前提下，Lever-LM 稳定优于 Random 及相似度检索。")

# ── 四、实验配置 ──
h("四、实验配置", 1)
table(["参数", "值"], [
    ["大模型", "Qwen2-VL-2B-Instruct"],
    ["Selector", "Autoregressive PointerSelector，d_model=1536"],
    ["K", "32"],
    ["shot_num", "2（sweep 时 1~8）"],
    ["reward_mode", "InfoScore"],
    ["训练损失", "RCE + neg_weight=1.0 InfoNCE"],
    ["VQA ckpt", "checkpoints/lever_lm_v2_poolneg/epoch003-valloss7.0768.ckpt"],
    ["ICL 评测", "VQAv2 validation 前 5000 条，unified 协议"],
])

# ── 五、实验结果 ──
h("五、实验结果", 1)

h("5.1 Selector 留出集（中间指标）", 2)
table(["任务", "排列对顺序 acc", "select 并集命中"], [
    ["VQA", "0.739", "0.154（≈随机 0.19）"],
    ["Caption", "0.667", "0.175"],
    ["SST-2", "0.447", "0.625"],
])
p("顺序问题已解决（>0.5）；VQA/Caption 上 top-32 内细选仍难。")

h("5.2 VQA 下游 ICL（unified 协议，val 5000，shot=2，K=32）", 2)
table(["方法", "2 万池", "44 万池", "说明"], [
    ["Lever-LM v2", "63.38%", "63.44%", "Selector ckpt 在 2 万池上训练"],
    ["Random", "61.70%", "63.16%", "扩池后 Random 明显上升"],
    ["0-shot", "70.54%", "70.54%", "无 ICD；诊断用，论文主表无此项"],
    ["Lever−Random", "+1.70%", "+0.28%", "与论文同向，幅度更小"],
])

h("5.3 Shot sweep Avg:1~8（2 万池，5000 val）", 2)
table(["方法", "Avg:1~8", "shot=1", "shot=2", "shot=8"], [
    ["Lever", "62.0%", "65.2%", "63.4%", "60.3%"],
    ["Random", "61.2%", "64.8%", "62.1%", "59.7%"],
    ["0-shot", "70.5%", "70.5%", "70.5%", "70.5%"],
])
p("Lever 全程略高于 Random；shot 越多 ICL 越低（负迁移）；0-shot 与 shot 无关。")

h("5.4 若不关注 0-shot，如何评价？", 2)
p(
    "仅看 ICL 内部（Lever vs Random）：结论方向与论文一致——Lever 优于 Random。"
    "但提升幅度小于原文（+0.3%~1.7% vs 论文 +3% 级）；"
    "扩 44 万池后 Lever 几乎不涨，说明 selector 需在全库分布上重训。"
)

# ── 六、核心问题：ICL 为何低于 0-shot？ ──
h("六、核心问题：ICL 为何低于 0-shot？", 1)
p(
    "2-shot 示例答案来自池中其他题的标注（如 pizza、donut），"
    "对示范自身可能正确，但对当前题（如 hot dog）是误导性例题。"
    "主问题答案不必、也不能从两个示范答案中二选一；模型自由生成，"
    "但常被 assistant 里的错示范带偏（约 20% 伤害 case 直接抄错 ICD 答案）。"
)

h("6.1 五层原因", 2)
table(["层级", "原因", "说明"], [
    ["检索", "按 (图,问) 相似度，不看答案", "top-32 里 ICD 答案=GT 仅 ~12~20%"],
    ["Prompt", "错答案写入 assistant 轮", "Chat 模型易模仿 → 抄 donut 而非 hot dog"],
    ["Selector", "训于 InfoScore，非下游 acc", "beam reward 打平，select 命中 ≈随机"],
    ["分布", "selector 在 2 万池训，44 万池推理", "扩池后 Random↑、Lever 不涨"],
    ["模型", "2B 0-shot 已 70.5%", "9B 论文 ICL 有更大提升空间；2B 怕噪声 ICD"],
])

h("6.2 2-shot 示例说明（非二选一）", 2)
p("当前题：小孩吃热狗，问 What is the kid eating?，GT=hot dog。")
p("示范1（池里另一题）：图=小孩吃 pizza，答 pizza —— 对示范自身正确，对当前题误导。")
p("示范2：图=小孩吃 donut，问句几乎相同，答 donut —— 模型最易抄 donut。")
p("0-shot 仅当前图+问，无 pizza/donut 干扰 → 更易答 hot dog。")

# ── 七、与论文差异小结 ──
h("七、与论文差异小结", 1)
table(["项目", "原版", "我们", "对齐？"], [
    ["5000 anchor", "train 随机 5000", "同左", "✅"],
    ["候选池规模", "全 train", "443,757 已建", "✅"],
    ["K / beam", "64 子支持 / beam 5", "K=32 / beam 3", "⚠️"],
    ["下游 metric", "VQA acc + CIDEr", "VQA 5000 val", "部分"],
    ["shot sweep", "Avg:1~8", "5000 val 已完成", "✅"],
    ["0-shot baseline", "主表无", "70.5%（诊断）", "N/A"],
    ["Lever>Random", "是（+3% 级）", "是（+0.3~1.7%）", "方向 ✅ 幅度 ⚠️"],
])

# ── 八、问题与解决方案 ──
h("八、问题与解决方案", 1)
table(["问题", "现状", "解决方案", "优先级"], [
    ["ICL < 0-shot", "63.4% vs 70.5%", "全库 Step1 重生成 + selector 重训；试 1-shot、K=64", "P0"],
    ["扩 44 万池 Lever 不涨", "+0.06%", "用 44 万池重跑 Step1（5000 anchor）+ 微调 selector", "P0"],
    ["Selector 信号弱", "select 命中≈随机", "InfoScore 直接选 ICD；增大 beam；池负样本已加", "P1"],
    ["Caption/SST-2 无下游", "未跑 CIDEr/acc", "扩展 icl_inference.py", "P2"],
    ["评测规模", "5000 子集", "可选扩至全 val 21.4 万", "P3"],
    ["与论文不可直接比", "2B vs 9B", "答辩强调工程创新 + ICL 内部 Lever>Random", "—"],
])

h("8.1 推荐下一步", 2)
bullet("P0：44 万池 + K=64 上，用全库重跑 Step1（5000 anchor，beam=5）→ 重训 selector → 再评 val5000")
bullet("P1：推理 baseline「InfoScore 贪心选 ICD」（绕开弱 selector 快速验证）")
bullet("P1：1-shot ICL（sweep 显示 65.2% > 2-shot 63.4%）")
bullet("P2：Caption CIDEr、SST-2 acc 下游 ICL")

# ── 九、结论 ──
h("九、结论", 1)
bullet("工程闭环完成：−∞ 根治、顺序可学、三任务数据+训练、VQA ICL+统一协议+sweep+44 万池。")
bullet("相对 v1：Teacher-forcing、InfoScore、解耦池、自回归 PointerSelector、InfoNCE 均有效。")
bullet("相对论文：ICL 内部 Lever>Random 方向一致，但幅度小；0-shot 高于 ICL 是 2B+强 0-shot 特有问题，论文主表未评 0-shot。")
bullet("核心矛盾：检索菜单常含「像但答案错」的 ICD → 2B 被带偏；需全库 Step1+重训 selector 释放 Lever 增益。")

p("")
p("附录：详细 Markdown 见 docs/项目总结_完整版_2026-06-05.md；"
  "44 万池 ICL 结果见 results/icl_fullpool_*_unified_n5000_k32.json。", bold=False)

out = "/home/jiyi/lizhiheng/Lever-LM/Lever-LM实验总结报告_最终版.docx"
doc.save(out)
print("Saved →", out)
